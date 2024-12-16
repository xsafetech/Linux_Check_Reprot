#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from bs4 import BeautifulSoup
import io
from PIL import Image, ImageDraw, ImageFont
from docx.oxml.ns import qn

def text_to_image(text, font_size=16, bg_color='black', text_color='white'):
    """
    将文本转换为终端风格的图片
    
    :param text: 要转换的文本
    :param font_size: 字体大小
    :param bg_color: 背景颜色
    :param text_color: 文本颜色
    :return: PIL Image对象
    """
    # 使用等宽字体,确保对齐，添加中文字体支持
    try:
        font_paths = [
            # 等宽中文字体优先
            #"/System/Library/Fonts/Menlo.ttc",  # macOS
            "/System/Library/Fonts/Menlo.ttc",
            #"C:/Windows/Fonts/sarasa-mono-sc-regular.ttf",  # Sarasa Mono SC
            #"C:/Windows/Fonts/SourceHanMonoSC-Regular.otf",  # Source Han Mono
            #"C:/Windows/Fonts/NotoSansMonoCJKsc-Regular.otf",  # Noto Sans Mono CJK
            #"/usr/share/fonts/sarasa-mono-sc-regular.ttf",  # Linux
            #"/usr/share/fonts/source-han-mono/SourceHanMonoSC-Regular.otf",
            # 备选等宽字体
            #"/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
            #"/System/Library/Fonts/Monaco.ttf",
            #"C:/Windows/Fonts/consola.ttf",
            #"consolas.ttf",
        ]
        font = None
        for path in font_paths:
            try:
                font = ImageFont.truetype(path, font_size)
                break
            except:
                continue
        if font is None:
            # 如果找不到合适的字体，使用默认等宽字体
            font = ImageFont.load_default()
    except:
        font = ImageFont.load_default()
    
    # 固定图片宽度和边距
    fixed_width = 1000  # 固定宽度
    padding_x = 40  # 水平内边距
    padding_y = 30  # 垂直内边距
    line_height = int(font_size * 1.5)  # 行高
    
    # 计算文本区域的最大宽度
    text_max_width = fixed_width - (padding_x * 2)
    
    # 处理文本换行
    lines = []
    for text_line in text.strip().split('\n'):
        # 如果单行超过最大宽度，保持原样（允许超出）
        lines.append(text_line)
    
    # 计算图片高度
    height = (len(lines) * line_height) + (padding_y * 2)
    
    # 创建高分辨率图片 (2x)
    scale = 2
    image = Image.new('RGB', (fixed_width * scale, height * scale), bg_color)
    draw = ImageDraw.Draw(image)
    
    # 缩放字体大小
    font_scaled = ImageFont.truetype(font.path, font_size * scale)
    
    # 绘制文本
    y = padding_y * scale
    for line in lines:
        draw.text((padding_x * scale, y), line, font=font_scaled, fill=text_color)
        y += line_height * scale
    
    # 缩放回原始大小
    image = image.resize((fixed_width, height), Image.Resampling.LANCZOS)
    
    return image


def convert_md_to_docx(md_file_path, output_dir=None):
    """
    将Markdown文件转换为Word文档
    
    :param md_file_path: Markdown文件路径
    :param output_dir: 输出目录，默认为Markdown文件所在目录
    """
    # 获取绝对路径
    md_file_path = os.path.abspath(md_file_path)
    
    # 检查文件是否存在
    if not os.path.exists(md_file_path):
        print(f"错误：文件 {md_file_path} 不存在")
        return None

    # 设置输出目录
    if output_dir is None:
        output_dir = os.path.dirname(md_file_path)
    else:
        output_dir = os.path.abspath(output_dir)
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 生成输出文件名
    base_name = os.path.splitext(os.path.basename(md_file_path))[0]
    docx_file_path = os.path.join(output_dir, f"{base_name}.docx")

    # 读取Markdown文件
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except Exception as e:
        print(f"读取文件错误：{e}")
        return None

    # 将Markdown转换为HTML
    html = markdown.markdown(md_content, extensions=['fenced_code', 'tables'])

    # 使用BeautifulSoup解析HTML
    soup = BeautifulSoup(html, 'html.parser')

    # 创建Word文档
    doc = Document()

    # 设置默认字体,也就是正文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10)
    # 设置中文字体，qn('w:eastAsia')Word文档中的东亚文字
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 创建一个字典来存储已处理过的文本内容，避免重复
    # 解析HTML并添加到Word文档
    # 修改解析逻辑，避免重复处理代码块
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'pre', 'code', ]):
        # 如果是嵌套在pre内的code，跳过处理
        if element.name == 'code' and element.parent.name == 'pre':
            continue
            
        if element.name.startswith('h'):
            heading = doc.add_heading(element.get_text(), level=int(element.name[1]))
            # 清除之前的样式
            heading.style._element.rPr.rFonts.clear()
            #获取 heading 的 run，并设置字体
            heading.style.font.bold = False
            # 设置标题加粗
            heading.style.font.bold = True
            # 设置标题颜色
            heading.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            heading.style.font.name = "宋体"
            heading.style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 根据标题级别设置不同的字体大小
            if int(element.name[1]) == 1:
                heading.style.font.size = Pt(22)    # 一级标题
            elif int(element.name[1]) == 2:
                heading.style.font.size = Pt(18)    # 二级标题
            elif int(element.name[1]) == 3:
                heading.style.font.size = Pt(16)    # 三级标题
            
        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                doc.add_paragraph(text)

        # 处理代码块（pre）和行内代码（独立的code），避免重复处理
        elif element.name in ['pre', 'code']:
            code_text = element.get_text().strip()
            if code_text:
                img = text_to_image(code_text)
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG', quality=95, dpi=(300, 300))
                img_byte_arr = img_byte_arr.getvalue()
                
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(io.BytesIO(img_byte_arr), width=Inches(7))

    # 保存Word文档
    doc.save(docx_file_path)
    print(f"转换完成：{docx_file_path}")
    return docx_file_path

def batch_convert(input_dir, output_dir=None):
    """
    批量转换目录中的所有Markdown文件
    
    :param input_dir: 输入目录
    :param output_dir: 输出目录，默认为输入目录
    """
    input_dir = os.path.abspath(input_dir)
    
    if output_dir is None:
        output_dir = input_dir
    else:
        output_dir = os.path.abspath(output_dir)

    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    # 遍历目录
    for filename in os.listdir(input_dir):
        if filename.endswith('.md'):
            md_path = os.path.join(input_dir, filename)
            convert_md_to_docx(md_path, output_dir)

def main():
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("用法：")
        print("  单个文件转换: python md_to_docx.py file.md")
        print("  批量转换目录: python md_to_docx.py /path/to/directory")
        sys.exit(1)

    input_path = sys.argv[1]
    
    # 获取绝对路径
    input_path = os.path.abspath(input_path)

    # 判断是文件还是目录
    if os.path.isfile(input_path):
        # 如果是相对路径文件，获取当前工作目录
        if not os.path.isabs(sys.argv[1]):
            input_path = os.path.join(os.getcwd(), sys.argv[1])
        convert_md_to_docx(input_path)
    elif os.path.isdir(input_path):
        batch_convert(input_path)
    else:
        print(f"错误：{input_path} 不是有效的文件或目录")

if __name__ == '__main__':
    main()